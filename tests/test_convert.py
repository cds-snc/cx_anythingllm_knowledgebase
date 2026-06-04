"""Tests for scripts/convert_documents.py.

Tests cover: slug generation, file hashing, manifest management,
new file conversion, idempotency, update detection, and bootstrapping
of existing manually-created markdown files.
"""

import json
import shutil
import textwrap
from pathlib import Path

import pytest

# Add scripts/ to path so we can import the module
import sys

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))

import convert_documents as cd


# ── Fixtures ─────────────────────────────────────────────────────────────────


@pytest.fixture
def workspace(tmp_path):
    """Create an isolated workspace with documents/ and docs/corpus-md/ dirs."""
    docs_dir = tmp_path / "documents"
    corpus_dir = tmp_path / "docs" / "corpus-md"
    docs_dir.mkdir()
    corpus_dir.mkdir(parents=True)

    # Monkeypatch module-level paths
    original_docs = cd.DOCUMENTS_DIR
    original_corpus = cd.CORPUS_DIR
    original_manifest = cd.MANIFEST_PATH

    cd.DOCUMENTS_DIR = docs_dir
    cd.CORPUS_DIR = corpus_dir
    cd.MANIFEST_PATH = corpus_dir / ".manifest.json"

    yield tmp_path, docs_dir, corpus_dir

    # Restore
    cd.DOCUMENTS_DIR = original_docs
    cd.CORPUS_DIR = original_corpus
    cd.MANIFEST_PATH = original_manifest


def _create_docx(path: Path, title: str, body: str) -> None:
    """Create a minimal DOCX file."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(body)
    doc.save(str(path))


def _create_xlsx(path: Path, headers: list, rows: list) -> None:
    """Create a minimal XLSX file."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(str(path))


# ── Unit Tests ───────────────────────────────────────────────────────────────


class TestSlugify:
    def test_simple_name(self):
        assert cd.slugify("my-document.pdf") == "my-document"

    def test_spaces_to_underscores(self):
        assert cd.slugify("My Document.docx") == "My_Document"

    def test_special_chars_removed(self):
        assert cd.slugify("Museum of Vancouver_ Guide.xlsx") == "Museum_of_Vancouver_Guide"

    def test_consecutive_underscores_collapsed(self):
        assert cd.slugify("file__with___many.pdf") == "file_with_many"

    def test_hyphens_preserved(self):
        assert cd.slugify("native-plants-for-shade.pdf") == "native-plants-for-shade"


class TestFileHash:
    def test_deterministic(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("hello world")
        h1 = cd.file_hash(f)
        h2 = cd.file_hash(f)
        assert h1 == h2

    def test_changes_with_content(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("version 1")
        h1 = cd.file_hash(f)
        f.write_text("version 2")
        h2 = cd.file_hash(f)
        assert h1 != h2

    def test_sha256_format(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("test")
        h = cd.file_hash(f)
        assert len(h) == 64  # SHA256 hex length
        assert all(c in "0123456789abcdef" for c in h)


class TestManifest:
    def test_load_empty(self, workspace):
        _, _, corpus_dir = workspace
        assert cd.load_manifest() == {}

    def test_save_and_load(self, workspace):
        _, _, corpus_dir = workspace
        data = {"file.pdf": {"hash": "abc123", "markdown": "file.md"}}
        cd.save_manifest(data)
        loaded = cd.load_manifest()
        assert loaded == data

    def test_sorted_keys(self, workspace):
        _, _, corpus_dir = workspace
        data = {"z.pdf": {"hash": "z", "markdown": "z.md"}, "a.pdf": {"hash": "a", "markdown": "a.md"}}
        cd.save_manifest(data)
        raw = (corpus_dir / ".manifest.json").read_text()
        keys = [k for k in json.loads(raw)]
        assert keys == ["a.pdf", "z.pdf"]


# ── Integration Tests ────────────────────────────────────────────────────────


class TestConvertNewFile:
    def test_converts_docx(self, workspace):
        _, docs_dir, corpus_dir = workspace
        _create_docx(docs_dir / "guide.docx", "Plant Guide", "Sword fern is great.")

        result = cd.main()
        assert result == 0

        md_path = corpus_dir / "guide.md"
        assert md_path.exists()
        content = md_path.read_text()
        assert "Plant Guide" in content
        assert "Sword fern" in content

    def test_converts_xlsx(self, workspace):
        _, docs_dir, corpus_dir = workspace
        _create_xlsx(
            docs_dir / "plants.xlsx",
            ["Name", "Type"],
            [["Sword Fern", "Shade"], ["Salal", "Sun"]],
        )

        result = cd.main()
        assert result == 0

        md_path = corpus_dir / "plants.md"
        assert md_path.exists()
        content = md_path.read_text()
        assert "Sword Fern" in content
        assert "Salal" in content

    def test_manifest_updated(self, workspace):
        _, docs_dir, corpus_dir = workspace
        _create_docx(docs_dir / "test.docx", "Test", "Content")

        cd.main()
        manifest = cd.load_manifest()
        assert "test.docx" in manifest
        assert "hash" in manifest["test.docx"]
        assert manifest["test.docx"]["markdown"] == "test.md"


class TestIdempotency:
    def test_second_run_skips(self, workspace, capsys):
        _, docs_dir, corpus_dir = workspace
        _create_docx(docs_dir / "test.docx", "Test", "Content")

        cd.main()
        first_content = (corpus_dir / "test.md").read_text()

        cd.main()
        captured = capsys.readouterr()
        assert "0 converted" in captured.out
        assert "1 unchanged" in captured.out

        # Content unchanged
        assert (corpus_dir / "test.md").read_text() == first_content


class TestUpdateDetection:
    def test_reconverts_changed_file(self, workspace, capsys):
        _, docs_dir, corpus_dir = workspace
        _create_docx(docs_dir / "guide.docx", "Version 1", "Old content")
        cd.main()

        old_content = (corpus_dir / "guide.md").read_text()
        assert "Version 1" in old_content

        # Modify the source
        _create_docx(docs_dir / "guide.docx", "Version 2", "New content")
        cd.main()

        new_content = (corpus_dir / "guide.md").read_text()
        assert "Version 2" in new_content
        assert "Version 1" not in new_content

        captured = capsys.readouterr()
        assert "1 converted" in captured.out


class TestBootstrapping:
    def test_preserves_existing_markdown(self, workspace, capsys):
        _, docs_dir, corpus_dir = workspace

        # Pre-create a markdown file (simulating manual conversion)
        (corpus_dir / "my-doc.md").write_text("# Manually Created\nHand-crafted content.")

        # Add corresponding source
        _create_docx(docs_dir / "my-doc.docx", "Auto Version", "Auto content")

        cd.main()
        captured = capsys.readouterr()
        assert "1 bootstrapped" in captured.out
        assert "0 converted" in captured.out

        # Original markdown preserved
        content = (corpus_dir / "my-doc.md").read_text()
        assert "Manually Created" in content
        assert "Hand-crafted" in content

    def test_bootstrap_then_unchanged(self, workspace, capsys):
        _, docs_dir, corpus_dir = workspace
        (corpus_dir / "doc.md").write_text("# Existing")
        _create_docx(docs_dir / "doc.docx", "Source", "Content")

        cd.main()  # bootstrap
        cd.main()  # should be unchanged

        captured = capsys.readouterr()
        assert "1 unchanged" in captured.out


class TestIgnoresUnsupported:
    def test_skips_txt_files(self, workspace, capsys):
        _, docs_dir, corpus_dir = workspace
        (docs_dir / "notes.txt").write_text("Some notes")

        cd.main()
        captured = capsys.readouterr()
        assert "0 converted" in captured.out
        assert not (corpus_dir / "notes.md").exists()

    def test_skips_directories(self, workspace, capsys):
        _, docs_dir, corpus_dir = workspace
        (docs_dir / "subdir").mkdir()

        cd.main()
        captured = capsys.readouterr()
        assert "0 converted" in captured.out


class TestFuzzyMatching:
    def test_underscore_vs_double_underscore(self, workspace):
        _, docs_dir, corpus_dir = workspace
        # Existing file with double underscore
        (corpus_dir / "My__Document.md").write_text("# Existing")
        _create_docx(docs_dir / "My_Document.docx", "Source", "Content")

        cd.main()
        manifest = cd.load_manifest()
        assert manifest["My_Document.docx"]["markdown"] == "My__Document.md"

    def test_truncated_filename(self, workspace):
        _, docs_dir, corpus_dir = workspace
        # Existing file with truncated name (>20 chars)
        (corpus_dir / "native-plants-by-flower-berry-colour-and-for-seasonal-intere.md").write_text(
            "# Existing"
        )
        _create_docx(
            docs_dir / "native-plants-by-flower-berry-colour-and-for-seasonal-interest.docx",
            "Source",
            "Content",
        )

        cd.main()
        manifest = cd.load_manifest()
        key = "native-plants-by-flower-berry-colour-and-for-seasonal-interest.docx"
        assert manifest[key]["markdown"] == "native-plants-by-flower-berry-colour-and-for-seasonal-intere.md"


class TestAutoSplit:
    def test_small_document_not_split(self, workspace):
        """A document below max_words should NOT be split."""
        _, docs_dir, corpus_dir = workspace
        # Create a DOCX with H2 headings but under 2000 words
        from docx import Document

        doc = Document()
        doc.add_heading("Main Guide", level=1)
        doc.add_paragraph("Introduction paragraph. " * 50)  # ~150 words
        doc.add_heading("Section 1", level=2)
        doc.add_paragraph("Content for section 1. " * 30)  # ~90 words
        doc.add_heading("Section 2", level=2)
        doc.add_paragraph("Content for section 2. " * 30)  # ~90 words
        doc.save(str(docs_dir / "small-guide.docx"))

        cd.main()
        # Should create a single file, not split
        assert (corpus_dir / "small-guide.md").exists()
        assert not (corpus_dir / "small-guide-Section_1.md").exists()

        manifest = cd.load_manifest()
        # Manifest should have string, not list
        assert isinstance(manifest["small-guide.docx"]["markdown"], str)

    def test_large_document_with_h2_gets_split(self, workspace):
        """A document over max_words with H2 headings should be split."""
        _, docs_dir, corpus_dir = workspace
        from docx import Document

        doc = Document()
        doc.add_heading("Plant Guide", level=1)
        doc.add_paragraph("Introduction. " * 400)  # ~400 words of intro

        # Add multiple H2 sections
        for i in range(1, 6):
            doc.add_heading(f"Plant Group {i}", level=2)
            doc.add_paragraph(f"Details about plant group {i}. " * 250)  # ~250 words each

        doc.save(str(docs_dir / "guide.docx"))

        cd.main()

        # Should create split files
        assert not (corpus_dir / "guide.md").exists()  # Original deleted
        assert (corpus_dir / "guide-Plant_Group_1.md").exists()
        assert (corpus_dir / "guide-Plant_Group_2.md").exists()
        assert (corpus_dir / "guide-Plant_Group_5.md").exists()

        manifest = cd.load_manifest()
        # Manifest should have list of files
        md_list = manifest["guide.docx"]["markdown"]
        assert isinstance(md_list, list)
        assert len(md_list) == 5
        assert "guide-Plant_Group_1.md" in md_list

    def test_split_preserves_preamble(self, workspace):
        """Preamble before first H2 should appear in all split files."""
        _, docs_dir, corpus_dir = workspace
        from docx import Document

        doc = Document()
        doc.add_heading("Guide Title", level=1)
        doc.add_paragraph("This is important preamble text that should appear in all sections.")

        doc.add_heading("Section A", level=2)
        doc.add_paragraph("Content A. " * 500)  # Larger to ensure > 2000 words

        doc.add_heading("Section B", level=2)
        doc.add_paragraph("Content B. " * 500)

        doc.save(str(docs_dir / "preamble-test.docx"))

        cd.main()

        # List what was created
        created_files = list(corpus_dir.glob("preamble-test-*.md"))

        # Read both split files and verify preamble is in both
        for md_file in created_files:
            content = md_file.read_text()
            assert "important preamble text" in content

    def test_split_files_no_content_lost(self, workspace):
        """Verify that splitting doesn't lose any content."""
        _, docs_dir, corpus_dir = workspace
        from docx import Document

        doc = Document()
        doc.add_heading("Original Guide", level=1)
        doc.add_paragraph("Intro content. " * 300)

        for i in range(1, 4):
            doc.add_heading(f"Section {i}", level=2)
            doc.add_paragraph(f"Unique content {i}. " * 250)

        doc.save(str(docs_dir / "content-test.docx"))

        cd.main()

        # Combine all split files
        all_content = ""
        for i in range(1, 4):
            md_file = corpus_dir / f"content-test-Section_{i}.md"
            assert md_file.exists()
            all_content += md_file.read_text()

        # Check that unique markers are present
        for i in range(1, 4):
            assert f"Unique content {i}" in all_content

    def test_document_without_h2_not_split(self, workspace):
        """Large document without H2 headings should NOT be split."""
        _, docs_dir, corpus_dir = workspace
        from docx import Document

        doc = Document()
        doc.add_heading("Big Guide", level=1)
        # Add lots of content without H2 structure
        for _ in range(5):
            doc.add_paragraph("Content paragraph. " * 350)

        doc.save(str(docs_dir / "no-structure.docx"))

        cd.main()

        # Should create single file, not split
        assert (corpus_dir / "no-structure.md").exists()
        assert not any(
            (corpus_dir / f).exists()
            for f in corpus_dir.glob("no-structure-*.md")
        )

        manifest = cd.load_manifest()
        assert isinstance(manifest["no-structure.docx"]["markdown"], str)

    def test_manifest_tracks_split_files(self, workspace):
        """Manifest should correctly track split files."""
        _, docs_dir, corpus_dir = workspace
        from docx import Document

        doc = Document()
        doc.add_heading("Test Guide", level=1)
        doc.add_paragraph("Preamble. " * 300)

        for section in ["Alpha", "Beta", "Gamma"]:
            doc.add_heading(section, level=2)
            doc.add_paragraph(f"Content for {section}. " * 300)

        doc.save(str(docs_dir / "manifest-test.docx"))

        cd.main()

        manifest = cd.load_manifest()
        assert "manifest-test.docx" in manifest
        assert "hash" in manifest["manifest-test.docx"]

        split_files = manifest["manifest-test.docx"]["markdown"]
        assert isinstance(split_files, list)
        assert len(split_files) == 3
        assert all(f.endswith(".md") for f in split_files)
        assert all((corpus_dir / f).exists() for f in split_files)
